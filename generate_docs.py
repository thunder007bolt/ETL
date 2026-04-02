"""
Génère deux fichiers Word de documentation pour le projet ETL CNSS :
  - documentation_technique.docx
  - documentation_fonctionnelle.docx

Usage :
    /home/thunder/anaconda3/bin/python generate_docs.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import copy

TODAY = datetime.date.today().strftime("%d/%m/%Y")
VERSION = "1.0"
ORG = "CNSS – Caisse Nationale de Sécurité Sociale"

# ─────────────────────────────────────────────────────────────────────────────
# Données GRH
# ─────────────────────────────────────────────────────────────────────────────

GRH_DIMS = [
    ("DTM.DIM_GRH_CAT_ABSENCE",          "Catégorie d'absence",               "CABS_CODE"),
    ("DTM.DIM_GRH_CAT_ARTICLE",          "Catégorie d'article de décision",   "CDEC_CODE, ORDRE"),
    ("DTM.DIM_GRH_CAT_DECISION",         "Catégorie de décision RH",          "CDEC_CODE"),
    ("DTM.DIM_GRH_CAT_DIPLOME",          "Catégorie de diplôme",              "CDIP_NO"),
    ("DTM.DIM_GRH_CAT_FRAIS",            "Catégorie de frais",                "CODE_CFR"),
    ("DTM.DIM_GRH_CAT_SANCTION",         "Catégorie de sanction disciplinaire","CSAN_NO"),
    ("DTM.DIM_GRH_CENTRE_MEDICAL",       "Centre médical",                    "CMED_CODE"),
    ("DTM.DIM_GRH_DOMAINE_ACTIVITE",     "Domaine d'activité",                "CODE_DOMAINE"),
    ("DTM.DIM_GRH_FONCTION",             "Fonction / poste",                  "FNCT_CODE"),
    ("DTM.DIM_GRH_FORMATION",            "Formation professionnelle",         "FORM_NO"),
    ("DTM.DIM_GRH_FUSION_ADMINISTRATIVE","Entité administrative fusionnée",   "CODE_ADM"),
    ("DTM.DIM_GRH_LIEU",                 "Lieu de travail / géographique",    "LIEU_ID"),
    ("DTM.DIM_GRH_NATIONALITE",          "Nationalité",                       "NATION_CODE"),
    ("DTM.DIM_GRH_NATURE_MOUVEMENT",     "Nature de mouvement RH",            "CODE_NAT_MVT"),
    ("DTM.DIM_GRH_PROFESSION",           "Profession",                        "PROF_CODE"),
    ("DTM.DIM_GRH_QUALIFICATION",        "Qualification / niveau",            "QUAL_CODE"),
    ("DTM.DIM_GRH_SECTION",              "Section organisationnelle",         "SECT_CODE"),
    ("DTM.DIM_GRH_SECTION_ANALYTIQUE",   "Section analytique (comptable)",    "CODE_ANA"),
    ("DTM.DIM_GRH_SENS_MOUVEMENT",       "Sens du mouvement (entrée/sortie)", "CODE_SENS_MVT"),
    ("DTM.DIM_GRH_TYPE_ENCADREMENT",     "Type d'encadrement hiérarchique",   "CODE_TYPE_ENCADRE"),
    ("DTM.DIM_GRH_TYPE_MOUVEMENT",       "Type de mouvement RH",              "CODE_TYPE_MVT"),
    ("DTM.DIM_GRH_UNITE_ADMIN_NATURE",   "Nature de l'unité administrative",  "UA_NATURE"),
    ("DTM.DIM_GRH_UNITE_ADMINISTRATIVE", "Unité administrative (organigramme)","UA_CODE"),
]

GRH_FACTS = [
    ("FAIT_GRH_PERSONNE",          "Snapshot mensuel de tous les agents (données personnelles, grade, contrat)"),
    ("FAIT_GRH_SITUATION",         "Situation administrative courante de chaque agent"),
    ("FAIT_GRH_SECTION",           "Affectation des agents par section organisationnelle"),
    ("FAIT_GRH_MOUVEMENT",         "Mouvements RH : recrutements, mutations, promotions, retraites"),
    ("FAIT_GRH_ABSENCE",           "Absences des agents (congés, maladies, etc.)"),
    ("FAIT_GRH_ABS_LIEU",          "Détail géographique des absences"),
    ("FAIT_GRH_DECISION",          "Décisions RH émises (arrêtés, notes de service)"),
    ("FAIT_GRH_DECISION_CARRIERE", "Décisions impactant la carrière de l'agent"),
    ("FAIT_GRH_DECISION_PERSONNE", "Lien décision ↔ agent concerné"),
    ("FAIT_GRH_DEMANDE",           "Demandes individuelles formulées par les agents"),
    ("FAIT_GRH_SANCTION",          "Sanctions disciplinaires appliquées"),
    ("FAIT_GRH_POSTE_FORMATIONS",  "Formations associées aux postes de travail"),
    ("FAIT_GRH_FACTURE",           "Factures liées aux frais RH (déplacements, formations)"),
]

GRH_DTM = [
    ("DTM.DTM_GRH_EFFECTIF",  "Effectif du personnel",
     "Répartition par tranche d'âge, ancienneté, sexe, domaine d'activité, lieu et fonction. "
     "Permet le suivi de la pyramide des âges et de l'évolution des effectifs."),
    ("DTM.DTM_GRH_RATIO",     "Ratios RH",
     "Indicateurs de ratio : genre, tranche d'âge, taux d'encadrement. "
     "Sert au pilotage de la diversité et de la structure managériale."),
    ("DTM.DTM_GRH_FORMATION", "Indicateurs de formation",
     "Heures de formation, coût, taux de réalisation par domaine et catégorie. "
     "Permet le suivi du plan de formation annuel."),
    ("DTM.DTM_GRH_MOUVEMENT", "Statistiques de mouvements",
     "Flux d'entrées et de sorties du personnel par type, nature et sens de mouvement. "
     "Sert à mesurer le turnover et la mobilité interne."),
]

# ─────────────────────────────────────────────────────────────────────────────
# Données Prestation / Recouvrement
# ─────────────────────────────────────────────────────────────────────────────

PR_DIMS = [
    ("DTM.DIM_TYPE_ACCIDENT",                "Type d'accident du travail",                "TAT_CODE"),
    ("DTM.DIM_TYPE_AJUSTEMENT",              "Type d'ajustement de cotisation",           "TAJ_CODE"),
    ("DTM.DIM_TYPE_ANOMALIE_BNTS",           "Type d'anomalie sur bordereau",             "TAB_CODE"),
    ("DTM.DIM_TYPE_BORDEREAU",               "Type de bordereau de versement",            "TBO_CODE"),
    ("DTM.DIM_TYPE_DEBOURS",                 "Type de débours",                           "CODE_TYPE_DEB"),
    ("DTM.DIM_TYPE_DOIGT",                   "Type de doigt (AT)",                        "CODE_DOIGT"),
    ("DTM.DIM_TYPE_DOSSIER",                 "Type de dossier prestation",                "TDOS_CODE"),
    ("DTM.DIM_TYPE_EFFET",                   "Type d'effet de paiement",                  "TEP_CODE"),
    ("DTM.DIM_TYPE_ETAPE",                   "Type d'étape de traitement dossier",        "TEA_CODE"),
    ("DTM.DIM_TYPE_FREQUENCE_PAIEMENT",      "Fréquence de paiement de prestation",       "TFP_CODE"),
    ("DTM.DIM_TYPE_MODE_PAIEMENT",           "Mode de paiement",                          "TMP_CODE"),
    ("DTM.DIM_TYPE_MOTIF",                   "Type de motif",                             "CODE_TYPE_MOTIF"),
    ("DTM.DIM_TYPE_OPER_CAISSE",             "Type d'opération de caisse",                "TOP_CODE"),
    ("DTM.DIM_TYPE_PIECE",                   "Type de pièce justificative",               "TPJ_CODE"),
    ("DTM.DIM_TYPE_PIECE_INIT_IND_ASSURE",   "Type de pièce initiale – assuré individuel","TPJ_CODE"),
    ("DTM.DIM_TYPE_PIECE_INIT_RECEPTION",    "Type de pièce initiale à la réception",     "TDOS_CODE, TPJ_CODE"),
    ("DTM.DIM_TYPE_PJ_CORRESPONDANCE",       "Type de pièce jointe – correspondance",     "TPJ_CODE"),
    ("DTM.DIM_TYPE_PRESTATION_ESP",          "Type de prestation en espèces",             "TPE_CODE"),
    ("DTM.DIM_TYPE_PRESTATION_NAT",          "Type de prestation en nature",              "TPN_CODE"),
    ("DTM.DIM_TYPE_PRESTATION_REGROUPE",     "Type de prestation regroupée",              "TPR_CODE"),
    ("DTM.DIM_TYPE_PRS_CORRESPONDANCE",      "Type de personne – correspondance",         "CODE_TYPE_PRS_CORRRESP"),
    ("DTM.DIM_BRANCHE",                      "Branche de sécurité sociale",               "BR_CODE"),
    ("DTM.DIM_BRANCHE_COTISATION",           "Branche de cotisation",                     "CODE_BRANCHE_COTISATION"),
    ("DTM.DIM_CAISSE_PAIEMENT",              "Caisse de paiement",                        "DR_NO, CAP_ID, LP_NO"),
    ("DTM.DIM_CAISSIER",                     "Agent caissier",                            "CAI_USERNAME"),
    ("DTM.DIM_CATEGORIE_DOSSIER_PENSION",    "Catégorie de dossier pension",              "DOS_CATEGORIE"),
    ("DTM.DIM_CATEGORIE_PRESTATION_NAT",     "Catégorie de prestation en nature",         "CPN_CODE"),
    ("DTM.DIM_COMPTE_BANCAIRE",              "Compte bancaire",                           "IF_NO, AG_CODE, COM_NO"),
    ("DTM.DIM_COMPTES",                      "Comptes de gestion",                        "CP_GES_NUMERO, CP_DIR_NUMERO, CP_NUMERO"),
    ("DTM.DIM_CONTROLEUR",                   "Contrôleur",                                "CON_ID"),
    ("DTM.DIM_DEPARTEMENT",                  "Département géographique",                  "PA_NO, PR_NO, DPT_NO"),
    ("DTM.DIM_DIRECTION_REGIONALE",          "Direction régionale CNSS",                  "DR_NO"),
    ("DTM.DIM_FORME_JURIDIQUE",              "Forme juridique de l'employeur",            "FJ_CODE"),
    ("DTM.DIM_LIEU_PAIEMENT",                "Lieu de paiement des prestations",          "LP_NO"),
    ("DTM.DIM_PARAM_CAISSE_REC_DEP",         "Paramètre caisse recettes/dépenses",        "CODE_CAISSE"),
    ("DTM.DIM_PARAM_CPT_BANCAIRE",           "Paramètre compte bancaire",                 "NUM_CPT_BANCAIRE"),
    ("DTM.DIM_PARAMETRES_COTISATION",        "Paramètres de cotisation",                  "PARC_JOURNAL"),
    ("DTM.DIM_PAYS",                         "Pays",                                      "PA_NO"),
    ("DTM.DIM_PERIODE",                      "Période de cotisation",                     "PER_ID"),
    ("DTM.DIM_AGENCE",                       "Agence bancaire",                           "IF_NO, AG_CODE"),
    ("DTM.DIM_ANNEE",                        "Année",                                     "AN_ID"),
    ("DTM.DIM_PERIODICITE_VERSEMENT",        "Périodicité de versement",                  "CODE_PERIODICITE"),
    ("DTM.DIM_PROVINCE",                     "Province géographique",                     "PR_NO, PA_NO"),
    ("DTM.DIM_REGIME_EMPLOYEUR",             "Régime de l'employeur",                     "EMP_REGIME"),
    ("DTM.DIM_SECTEUR_ACTIVITE",             "Secteur d'activité économique",             "SA_NO"),
    ("DTM.DIM_SECTEUR_OPERATION",            "Secteur d'opération de l'employeur",        "EMP_ID, SA_NO, SSA_NO, SO_DATE_DEBUT"),
    ("DTM.DIM_SEMESTRE",                     "Semestre",                                  "SEM_ID"),
    ("DTM.DIM_SERVICE_PROVINCIAL",           "Service provincial",                        "SP_NO"),
    ("DTM.DIM_REGION",                       "Région (direction, service, lieu)",         "DR_NO, SP_NO, LP_NO"),
    ("DTM.DIM_TRANCHE_AGE",                  "Tranche d'âge",                             "TAG_CODE"),
    ("DTM.DIM_TRANCHE_EFFECTIF",             "Tranche d'effectif",                        "TEF_CODE"),
    ("DTM.DIM_TRIMESTRE",                    "Trimestre",                                 "TRIM_ID"),
]

PR_FACTS = [
    ("FAIT_CATEGORIE_EMPLOYEUR",     "Catégorie d'employeur par période"),
    ("FAIT_CONTRAINTE",              "Contraintes juridiques imposées aux employeurs"),
    ("FAIT_CONTROLE",                "Contrôles effectués chez les employeurs"),
    ("FAIT_DECLARATION_NOMINATIVE",  "Déclarations nominatives des salariés"),
    ("FAIT_DEPOT",                   "Dépôts de bordereaux de cotisation"),
    ("FAIT_DOSSIER_IMMATRICULATION", "Dossiers d'immatriculation employeur/travailleur"),
    ("FAIT_EMPLOI",                  "Emplois déclarés par les employeurs"),
    ("FAIT_EMPLOYEUR",               "Snapshot mensuel des employeurs affiliés"),
    ("FAIT_INDIVIDU",                "Individus assurés sociaux"),
    ("FAIT_MISE_EN_DEMEURE",         "Mises en demeure pour non-paiement de cotisations"),
    ("FAIT_NOTIFICATION_PREST",      "Notifications de prestations"),
    ("FAIT_PARTENAIRE_ASSURANCE",    "Partenaires assurance groupe"),
    ("FAIT_PENALITE_RETARD",         "Pénalités de retard sur cotisations"),
    ("FAIT_PRESTATION_ESP",          "Prestations en espèces versées"),
    ("FAIT_RECEPTION_DOSSIER",       "Réceptions de dossiers de prestation"),
    ("FAIT_TRANSACTION_COTISATION",  "Transactions de cotisation (appels, règlements)"),
    ("FAIT_TRANSACTION_DECLARATION", "Transactions liées aux déclarations"),
    ("FAIT_TRANSACTION_REGLEMENT",   "Transactions de règlement des cotisations"),
    ("FAIT_TRAVAILLEUR",             "Travailleurs déclarés"),
    ("FAIT_TSINISTRE",               "Sinistres (accidents du travail, maladies)"),
    ("FAIT_APPEL",                   "Appels de cotisation émis"),
    ("FAIT_BORDEREAU_PREST",         "Bordereaux de prestation"),
    ("FAIT_DEBOURS",                 "Débours effectués"),
    ("FAIT_DECLARAT_GROUPE_ASSURANCE","Déclarations groupe assurance"),
    ("FAIT_DOSSIER",                 "Dossiers de prestation (détail)"),
    ("FAIT_ECHEANCIER",              "Échéanciers de paiement accordés"),
    ("FAIT_EFFET",                   "Effets de paiement"),
    ("FAIT_ELEMENT_CONTROLE",        "Éléments de contrôle"),
    ("FAIT_ELEMENT_TAXATION",        "Éléments de taxation d'office"),
    ("FAIT_ETAPE",                   "Étapes de traitement des dossiers"),
    ("FAIT_IND_COMPTE_BANCAIRE",     "Comptes bancaires des individus assurés"),
    ("FAIT_INSTITUTION_FINANCIERE",  "Institutions financières partenaires"),
    ("FAIT_LIEN",                    "Liens familiaux (bénéficiaires)"),
    ("FAIT_PERIODE_COTISATION",      "Périodes de cotisation par employeur"),
    ("FAIT_AJUSTEMENT",              "Ajustements de cotisation"),
    ("FAIT_PE_AJ",                   "Périodes d'ajustement"),
    ("FAIT_RECEPTION_DOS_PIECE",     "Pièces jointes reçues avec les dossiers"),
    ("FAIT_REDEVANCE",               "Redevances"),
    ("FAIT_SALAIRE",                 "Salaires déclarés par les employeurs"),
    ("FAIT_SUPPLEMENT_SALAIRE",      "Suppléments de salaire"),
    ("FAIT_TAXATION_DOFFICE",        "Taxations d'office"),
    ("FAIT_TXDEPOT",                 "Taux de dépôt des déclarations"),
]

PR_DTM = [
    ("DTM.DTM_EMPLOYEUR",         "Indicateurs Employeurs",
     "Suivi de la population employeurs : effectif, régime, forme juridique, secteur d'activité."),
    ("DTM.DTM_COTISATIONS",       "Cotisations (AT + Vieillesse)",
     "Montants appelés, encaissés et reversés par branche (accidents du travail, assurance vieillesse). "
     "Calcul du taux de recouvrement par période."),
    ("DTM.DTM_RECOUVREMENT",      "Indicateurs de Recouvrement",
     "Taux de recouvrement global, montants en retard, historique de règlement."),
    ("DTM.DTM_CONTROLE",          "Statistiques de Contrôle",
     "Nombre de contrôles réalisés, résultats, redressements, par direction régionale."),
    ("DTM.DTM_SALAIRE",           "Analyse des Salaires",
     "Masse salariale déclarée par branche, tranche d'effectif et secteur d'activité."),
    ("DTM.DTM_MISE_EN_DEMEURE",   "Mises en Demeure",
     "Suivi des mises en demeure émises, réglées et en cours par direction et secteur."),
    ("DTM.DTM_TRAVAILLEUR",       "Travailleurs Déclarés",
     "Nombre de travailleurs déclarés par employeur, branche et période."),
    ("DTM.DTM_IMM_EMPLOYEUR",     "Immatriculation Employeurs",
     "Flux de nouvelles immatriculations d'employeurs par période et direction."),
    ("DTM.DTM_IMM_TRAVAILLEUR",   "Immatriculation Travailleurs",
     "Flux de nouvelles immatriculations de travailleurs."),
    ("DTM.DTM_PRESTATION",        "Prestations Versées",
     "Montant et nombre de prestations en espèces et en nature versées par type et branche."),
    ("DTM.DTM_PRESTATION_INDUE",  "Prestations Indues",
     "Prestations versées à tort : montants récupérés, en cours de récupération."),
    ("DTM.DTM_ACCIDENT_TRAVAIL",  "Accidents du Travail",
     "Fréquence, gravité et coût des accidents du travail par secteur d'activité."),
    ("DTM.DTM_DOSSIER",           "Dossiers de Prestation",
     "Suivi du volume et de l'état des dossiers par type, étape et direction."),
]


# ─────────────────────────────────────────────────────────────────────────────
# Helpers de mise en forme
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Applique une couleur de fond à une cellule de tableau Word."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_header_row(table, headers, bg="1F3864"):
    """Ajoute une ligne d'en-tête colorée à un tableau."""
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        cell.text = hdr
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table_rows(table, rows, start=1):
    """Remplit les lignes d'un tableau (alternance de couleurs)."""
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8.5)
            if idx % 2 == 0:
                set_cell_bg(cell, "EBF1F9")


def style_table(table):
    """Style global du tableau."""
    table.style = 'Table Grid'
    for col in table.columns:
        for cell in col.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)


def add_code_block(doc, text):
    """Ajoute un bloc de code (style Courier, fond gris)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6C)
    # fond gris clair sur le paragraphe
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p


def add_cover(doc, title, subtitle):
    """Page de garde."""
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ORG)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version : {VERSION}     |     Date : {TODAY}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_page_break()


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def body(doc, text):
    return doc.add_paragraph(text)


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Document 1 — TECHNIQUE
# ─────────────────────────────────────────────────────────────────────────────

def build_technical_doc():
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    add_cover(doc,
              "Documentation Technique",
              "Système ETL – Data Warehouse CNSS")

    # ── 1. Introduction ──────────────────────────────────────────────────────
    h1(doc, "1. Introduction")
    body(doc,
         "Ce document décrit l'architecture technique, les composants Python et la logique de "
         "traitement du système ETL (Extract-Transform-Load) du Data Warehouse de la CNSS. "
         "Il est destiné aux développeurs et administrateurs chargés du déploiement, de la "
         "maintenance et de l'évolution de la plateforme.")
    body(doc,
         "Le système couvre deux domaines métier :")
    bullet(doc, "GRH (Gestion des Ressources Humaines)")
    bullet(doc, "Prestation / Recouvrement")

    # ── 2. Architecture globale ──────────────────────────────────────────────
    h1(doc, "2. Architecture Globale")
    body(doc,
         "L'ETL suit une architecture trois couches, du système source OLTP jusqu'aux tables "
         "agrégées consommées par les outils de Business Intelligence :")

    add_code_block(doc,
        "Source OLTP (Oracle)\n"
        "      ↓  [Extract]\n"
        "Staging Parquet  (staging/raw/ → staging/transformed/)\n"
        "      ↓  [Transform]\n"
        "Data Warehouse Oracle\n"
        "   ├─ DWH.FAIT_*   (tables de faits – données transactionnelles détaillées)\n"
        "   └─ DWH.DIM_*    (tables de dimensions – données de référence)\n"
        "      ↓  [Load DTM]\n"
        "Couche DTM Oracle\n"
        "   └─ DTM.DTM_*    (agrégats pré-calculés pour le BI)\n"
        "      └─ DTM.DIM_* (dimensions partagées)")

    h2(doc, "2.1 Modes d'exécution")
    body(doc,
         "Chaque pipeline peut être exécuté en mode granulaire (E / T / L) "
         "ou en mode ETL complet en mémoire (sans étape intermédiaire) :")

    table = doc.add_table(rows=1, cols=3)
    style_table(table)
    add_header_row(table, ["Mode", "Description", "Stockage intermédiaire"])
    add_table_rows(table, [
        ("E – Extract",   "Extraction depuis la base source → fichier Parquet",       "staging/raw/{nom}_{YYYYMMDD}.parquet"),
        ("T – Transform", "Lecture Parquet raw, renommage, enrichissement → Parquet",  "staging/transformed/{nom}_{YYYYMMDD}.parquet"),
        ("L – Load",      "Lecture Parquet transformed, chargement dans Oracle DWH",   "–"),
        ("ETL complet",   "E+T+L enchaînés en mémoire, sans fichier intermédiaire",    "Aucun"),
    ])

    # ── 3. Structure du projet ───────────────────────────────────────────────
    h1(doc, "3. Structure du Projet")
    add_code_block(doc,
        "ETL/\n"
        "├── domaines/\n"
        "│   ├── grh/\n"
        "│   │   ├── dims/        ← pipeline + config + SQL des dimensions GRH\n"
        "│   │   ├── facts/       ← pipeline + config + SQL des faits GRH\n"
        "│   │   └── dtm/         ← pipeline + config + SQL des agrégats GRH\n"
        "│   └── prestation_recouvrement/\n"
        "│       ├── dims/        ← pipeline + config + SQL des dimensions P/R\n"
        "│       ├── facts/       ← pipeline + config + SQL des faits P/R\n"
        "│       └── dtm/         ← pipeline + config + SQL des agrégats P/R\n"
        "├── shared/\n"
        "│   ├── base/            ← classes abstraites BaseLoader, BaseExtractor\n"
        "│   ├── configs/         ← settings.py, log_setup.py\n"
        "│   └── utils/           ← db_utils.py, staging.py, sql_loader.py\n"
        "├── run_dim.py           ← CLI – chargement d'une dimension\n"
        "├── run_fact.py          ← CLI – chargement d'un fait\n"
        "├── run_dtm.py           ← CLI – chargement des agrégats DTM\n"
        "├── requirements.txt\n"
        "└── .env                 ← variables d'environnement (credentials)")

    # ── 4. Environnement et Configuration ───────────────────────────────────
    h1(doc, "4. Environnement et Configuration")

    h2(doc, "4.1 Fichier .env")
    body(doc,
         "Les paramètres de connexion et les chemins sont définis dans le fichier .env "
         "à la racine du projet (chargé automatiquement par python-dotenv) :")
    add_code_block(doc,
        "# Base de données source (OLTP)\n"
        "SRC_ORACLE_HOST=<adresse_serveur>\n"
        "SRC_ORACLE_PORT=1521\n"
        "SRC_ORACLE_SERVICE=<service_oracle>\n"
        "SRC_ORACLE_USER=user_dwh\n"
        "SRC_ORACLE_PASSWORD=<mot_de_passe>\n\n"
        "# Base Data Warehouse (cible)\n"
        "DW_ORACLE_HOST=<adresse_serveur>\n"
        "DW_ORACLE_PORT=1521\n"
        "DW_ORACLE_SERVICE=<service_oracle>\n"
        "DW_ORACLE_USER=dwh\n"
        "DW_ORACLE_PASSWORD=<mot_de_passe>\n\n"
        "# Staging et archivage\n"
        "STAGING_DIR=/chemin/vers/staging\n"
        "ODS_SCHEMA=ods          # optionnel – active l'archivage ODS avant TRUNCATE\n"
        "ORA_THICK_MODE=false     # true si Oracle Instant Client requis")

    h2(doc, "4.2 Dépendances Python")
    body(doc, "Fichier requirements.txt :")
    add_code_block(doc,
        "oracledb       # Pilote Oracle (thin ou thick mode)\n"
        "pandas         # Manipulation de DataFrames\n"
        "numpy          # Opérations numériques\n"
        "python-dotenv  # Chargement du fichier .env\n"
        "pyyaml         # Fichiers de configuration YAML\n"
        "pyarrow        # Format Parquet (staging intermédiaire)\n"
        "colorlog       # Logs colorés dans le terminal")

    # ── 5. Classes de base Python ────────────────────────────────────────────
    h1(doc, "5. Classes de Base Python")

    h2(doc, "5.1 BaseLoader (shared/base/base_loader.py)")
    body(doc,
         "Classe abstraite centralisant toutes les opérations d'écriture vers Oracle. "
         "Tous les loaders de dimension et de fait héritent de cette classe.")
    table = doc.add_table(rows=1, cols=2)
    style_table(table)
    add_header_row(table, ["Méthode", "Description"])
    add_table_rows(table, [
        ("_merge(table, df, key_cols, seq_cols)",
         "MERGE Oracle : mise à jour des lignes existantes, insertion des nouvelles. "
         "Utilise les key_cols comme clés de jointure."),
        ("_merge_via_gtt(table, df, key_cols, seq_cols)",
         "Variante GTT : passe par une table temporaire globale pour les grandes dimensions. "
         "Réduit la pression Redo/Undo."),
        ("_bulk_insert(table, df, seq_cols)",
         "INSERT /*+ APPEND */ par lots. Utilisé pour les tables de faits volumineuses."),
        ("_delete_cliche(table, cliche)",
         "DELETE WHERE CLICHE = :1. Supprime les données du snapshot mensuel avant rechargement."),
        ("_archive_to_ods_and_truncate(table, ods_schema, cliche)",
         "Archive les données dans le schéma ODS puis TRUNCATE la table DWH."),
        ("_delete_insert_period(table, df, period_cols, seq_cols)",
         "Delete par période + INSERT pour les tables avec colonnes L_ANNEE / L_MOIS."),
        ("_ensure_gtt(table)",
         "Crée la table temporaire globale (GTT) si elle n'existe pas déjà."),
    ])

    h2(doc, "5.2 BaseExtractor (shared/base/base_extractor.py)")
    body(doc,
         "Classe abstraite pour l'extraction depuis Oracle. Gère la connexion, "
         "le streaming par lots (fetchmany), et l'écriture Parquet incrémentale.")

    h2(doc, "5.3 BaseTransformer (shared/base/base_transformer.py)")
    body(doc,
         "Classe abstraite pour les transformations de données : renommage des colonnes "
         "(col_map), ajout de colonnes synthétiques (extra_cols), et application "
         "de fonctions de transformation métier (transform_fn).")

    # ── 6. Pipelines ─────────────────────────────────────────────────────────
    h1(doc, "6. Pipelines ETL")

    h2(doc, "6.1 Pipeline Dimensions (pipeline_dims.py)")
    body(doc,
         "Chaque domaine possède son propre pipeline de dimensions. "
         "La classe DimsPipeline lit la configuration (dim_config.py), "
         "établit les connexions source et DWH, puis charge chaque dimension "
         "via la classe _GenericDimLoader.")
    body(doc, "Stratégies de chargement disponibles :")
    bullet(doc, "merge (défaut) : MERGE Oracle – upsert ligne par ligne via les key_cols.")
    bullet(doc, "gtt : MERGE via GTT – recommandé pour les dimensions volumineuses (> 100 000 lignes).")

    h2(doc, "6.2 Pipeline Faits (pipeline_facts.py)")
    body(doc,
         "La classe FactsPipeline orchestre le chargement des tables de faits. "
         "L'extraction est réalisée en streaming (fetchmany) pour éviter les problèmes "
         "mémoire sur les tables volumineuses. Le CLICHE est injecté automatiquement "
         "dans chaque lot à partir de la date du run.")
    body(doc, "Séquence de chargement d'un fait :")
    bullet(doc, "1. Extraction SQL depuis la base source (avec cursor.arraysize configurable)")
    bullet(doc, "2. Conversion en DataFrame Pandas avec cast des types Oracle")
    bullet(doc, "3. Injection du CLICHE (format MMYYYY)")
    bullet(doc, "4. Suppression des données existantes pour ce CLICHE (ou archive ODS)")
    bullet(doc, "5. INSERT /*+ APPEND */ par lots dans la table de fait DWH")

    h2(doc, "6.3 Pipeline DTM (pipeline_dtm.py)")
    body(doc,
         "Le pipeline DTM exécute les requêtes SQL d'agrégation qui lisent les tables "
         "de faits DWH et les dimensions DTM. Le résultat est chargé dans les tables "
         "DTM.DTM_* après suppression des données du CLICHE courant.")
    body(doc,
         "Les requêtes SQL DTM reçoivent :1 = CLICHE comme variable de liaison "
         "pour filtrer le snapshot mensuel source.")

    # ── 7. Scripts CLI ───────────────────────────────────────────────────────
    h1(doc, "7. Scripts d'Exécution (CLI)")

    h2(doc, "7.1 run_dim.py – Chargement d'une Dimension")
    table = doc.add_table(rows=1, cols=3)
    style_table(table)
    add_header_row(table, ["Argument", "Valeurs", "Description"])
    add_table_rows(table, [
        ("--dim",  "<nom_dimension>",         "Nom court de la dimension (ex : branche, grh_fonction)"),
        ("--step", "E | T | L | (absent)",    "Étape à exécuter. Sans --step : ETL complet en mémoire"),
    ])
    body(doc, "Exemples :")
    add_code_block(doc,
        "# ETL complet en mémoire\n"
        "python run_dim.py --dim branche\n\n"
        "# Extraction seule vers staging/raw/\n"
        "python run_dim.py --dim branche --step E\n\n"
        "# Transformation seule\n"
        "python run_dim.py --dim branche --step T\n\n"
        "# Chargement depuis staging/transformed/\n"
        "python run_dim.py --dim branche --step L")

    h2(doc, "7.2 run_fact.py – Chargement d'un Fait")
    table = doc.add_table(rows=1, cols=3)
    style_table(table)
    add_header_row(table, ["Argument", "Valeurs", "Description"])
    add_table_rows(table, [
        ("--fact",  "<nom_fait>",            "Nom court du fait (ex : contrainte, grh_absence)"),
        ("--step",  "E | T | L | (absent)",  "Étape à exécuter. Sans --step : ETL complet en mémoire"),
        ("--date",  "YYYYMMDD",              "Date de référence pour les fichiers staging (défaut : aujourd'hui)"),
        ("--fetch", "entier",                "Taille des lots d'extraction (défaut : 50 000)"),
    ])
    body(doc, "Exemples :")
    add_code_block(doc,
        "# ETL complet\n"
        "python run_fact.py --fact contrainte\n\n"
        "# Extraction avec lot de 100 000 lignes\n"
        "python run_fact.py --fact grh_absence --step E --fetch 100000\n\n"
        "# Chargement sur une date précise\n"
        "python run_fact.py --fact contrainte --step L --date 20260301")

    h2(doc, "7.3 run_dtm.py – Chargement des Agrégats DTM")
    table = doc.add_table(rows=1, cols=3)
    style_table(table)
    add_header_row(table, ["Argument", "Valeurs", "Description"])
    add_table_rows(table, [
        ("--dtm",   "<noms séparés par espace> | (absent)", "Tables DTM à charger (défaut : toutes)"),
        ("--date",  "YYYYMMDD",                             "Date de référence pour le CLICHE"),
        ("--fetch", "entier",                               "Taille des lots (défaut : 100 000)"),
    ])
    body(doc, "Exemples :")
    add_code_block(doc,
        "# Toutes les tables DTM\n"
        "python run_dtm.py\n\n"
        "# Deux tables spécifiques\n"
        "python run_dtm.py --dtm cotisations recouvrement\n\n"
        "# Sur une date précise\n"
        "python run_dtm.py --dtm grh_effectif --date 20260301\n\n"
        "# Taille des lots personnalisée\n"
        "python run_dtm.py --fetch 10000")

    # ── 8. Identifiant CLICHE ────────────────────────────────────────────────
    h1(doc, "8. Identifiant de Snapshot (CLICHE)")
    body(doc,
         "Le CLICHE est l'identifiant de snapshot mensuel utilisé pour tracer et "
         "partitionner les données dans les tables de faits et DTM.")
    table = doc.add_table(rows=1, cols=2)
    style_table(table)
    add_header_row(table, ["Propriété", "Valeur"])
    add_table_rows(table, [
        ("Format",       "MMYYYY  (ex : 032026 pour Mars 2026)"),
        ("Source",       "Calculé automatiquement à partir de la date du run"),
        ("Usage DWH",    "Colonne CLICHE dans toutes les tables de faits"),
        ("Usage DTM",    "Variable de liaison :1 dans les SQL d'agrégation"),
        ("Rechargement", "DELETE WHERE CLICHE = :1 avant INSERT – replay idempotent"),
    ])
    body(doc,
         "Cette conception permet de rejouer un mois entier sans effacer les autres "
         "snapshots, et facilite les corrections rétroactives.")

    # ── 9. Stratégies de Chargement ─────────────────────────────────────────
    h1(doc, "9. Stratégies de Chargement")

    h2(doc, "9.1 MERGE – Dimensions")
    body(doc,
         "Stratégie par défaut pour les dimensions. Utilise l'instruction MERGE Oracle "
         "sur les colonnes clés (key_cols). Met à jour les lignes existantes et insère "
         "les nouvelles. Les clés de substitution (seq_cols) sont générées via des "
         "séquences Oracle uniquement à l'insertion.")

    h2(doc, "9.2 MERGE via GTT – Grandes Dimensions")
    body(doc,
         "Pour les dimensions volumineuses (> 100 000 lignes), une table temporaire "
         "globale (GTT) est utilisée comme intermédiaire. Les données sont d'abord "
         "chargées dans la GTT puis fusionnées dans la table cible. "
         "Cette approche réduit significativement la pression sur les segments "
         "Redo et Undo d'Oracle.")

    h2(doc, "9.3 Delete + Insert – Tables de Faits")
    body(doc,
         "Stratégie standard pour les tables de faits :")
    bullet(doc, "Suppression des données du snapshot courant : DELETE WHERE CLICHE = :1")
    bullet(doc, "Insertion des nouvelles données : INSERT /*+ APPEND */")
    body(doc,
         "Cette approche garantit l'idempotence : un mois peut être rechargé autant "
         "de fois que nécessaire sans risque de doublon.")

    h2(doc, "9.4 Archive ODS + Truncate")
    body(doc,
         "Stratégie alternative activée lorsque ODS_SCHEMA est défini dans .env. "
         "Utilisée pour les très grandes tables de faits :")
    bullet(doc, "1. Copie des données dans le schéma ODS (archive longue durée)")
    bullet(doc, "2. TRUNCATE de la table DWH (nettoyage complet)")
    bullet(doc, "3. INSERT des nouvelles données")

    # ── 10. Gestion des Types Oracle / Pandas ───────────────────────────────
    h1(doc, "10. Gestion des Types Oracle / Pandas")
    body(doc,
         "La conversion entre les types Oracle et Pandas est critique pour éviter "
         "les erreurs d'insertion. Le pipeline applique les règles suivantes :")
    table = doc.add_table(rows=1, cols=3)
    style_table(table)
    add_header_row(table, ["Type Oracle", "Type Pandas résultant", "Note"])
    add_table_rows(table, [
        ("NUMBER(p,0)",  "Int64 (nullable)",      "Entier avec NULL → utilise pd.Int64Dtype()"),
        ("NUMBER(p,s>0)","float64",               "Décimal"),
        ("DATE / TIMESTAMP","datetime64[ns]",     "Conversion automatique"),
        ("VARCHAR2 / CHAR","object (str)",        "Chaîne de caractères"),
        ("NULL Oracle",  "pd.NA / pd.NaT / NaN",  "Géré avant l'INSERT Oracle"),
        ("decimal.Decimal","float64",             "Converti pour compatibilité Pandas"),
    ])
    body(doc,
         "La cohérence de schéma entre les lots (batches) est assurée par un cast "
         "explicite vers le schéma du premier batch, évitant les dérives de type "
         "lorsque des NULL apparaissent dans un lot ultérieur.")

    # ── 11. Staging Parquet ──────────────────────────────────────────────────
    h1(doc, "11. Staging Parquet")
    body(doc,
         "Le staging Parquet est utilisé en mode ETL granulaire (étapes E / T / L séparées). "
         "Les fichiers sont nommés selon la convention suivante :")
    add_code_block(doc,
        "staging/\n"
        "├── raw/          ← fichiers extraits de la source\n"
        "│   └── {nom}_{YYYYMMDD}.parquet\n"
        "└── transformed/  ← fichiers après transformation\n"
        "    └── {nom}_{YYYYMMDD}.parquet")
    body(doc,
         "L'écriture est incrémentale (PyArrow ParquetWriter) : les données sont "
         "écrites lot par lot sans charger l'intégralité en mémoire. "
         "Le schéma est figé sur le premier lot pour garantir la cohérence.")

    # ── 12. Logging ──────────────────────────────────────────────────────────
    h1(doc, "12. Logging et Observabilité")
    body(doc,
         "Le module shared/configs/log_setup.py configure le système de logs :")
    table = doc.add_table(rows=1, cols=2)
    style_table(table)
    add_header_row(table, ["Niveau", "Usage"])
    add_table_rows(table, [
        ("DEBUG",    "Détails internes (requêtes SQL, compteurs de lots)"),
        ("INFO",     "Progression normale : début/fin de chaque table, compteur de lignes"),
        ("WARNING",  "Situations inattendues mais non bloquantes (ex : table vide)"),
        ("ERROR",    "Erreurs récupérées (ex : perte de connexion transiente)"),
        ("CRITICAL", "Erreurs fatales entraînant l'arrêt du processus"),
    ])
    body(doc,
         "Les logs sont écrits simultanément vers la console (avec couleurs via colorlog) "
         "et vers un fichier dans le répertoire logs/.")

    # ── 13. Sécurité ─────────────────────────────────────────────────────────
    h1(doc, "13. Sécurité et Bonnes Pratiques")
    bullet(doc, "Les credentials de base de données sont stockés uniquement dans .env, "
                "jamais dans le code source ni dans le dépôt git (.gitignore).")
    bullet(doc, "Le compte source (user_dwh) dispose uniquement des droits SELECT "
                "sur les schémas métier. Aucun droit d'écriture sur la source.")
    bullet(doc, "Toutes les requêtes paramétrées utilisent des bind variables Oracle "
                "(:1, :2, …) pour prévenir les injections SQL.")
    bullet(doc, "Le mode GTT et le streaming évitent de charger des jeux de données "
                "complets en mémoire, réduisant les risques d'exposition de données sensibles.")

    doc.save("documentation_technique.docx")
    print("✔  documentation_technique.docx généré")


# ─────────────────────────────────────────────────────────────────────────────
# Document 2 — FONCTIONNEL
# ─────────────────────────────────────────────────────────────────────────────

def build_functional_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    add_cover(doc,
              "Documentation Fonctionnelle",
              "Système ETL – Data Warehouse CNSS")

    # ── 1. Présentation ───────────────────────────────────────────────────────
    h1(doc, "1. Présentation de la CNSS et Contexte du Projet")
    body(doc,
         "La Caisse Nationale de Sécurité Sociale (CNSS) est l'organisme tunisien chargé "
         "de la gestion de la sécurité sociale pour les salariés du secteur privé. "
         "Elle administre notamment les branches assurance vieillesse, accidents du travail, "
         "prestations familiales et assurance maladie groupe.")
    body(doc,
         "Le projet Data Warehouse ETL a pour objectif de centraliser et d'agréger "
         "les données métier issues du système opérationnel (OLTP) afin de produire "
         "des indicateurs de pilotage fiables et actualisés mensuellement.")
    body(doc,
         "Deux domaines métier sont couverts par ce système :")
    bullet(doc, "GRH (Gestion des Ressources Humaines) : suivi du personnel interne de la CNSS.")
    bullet(doc,
           "Prestation / Recouvrement : gestion des affiliés (employeurs et travailleurs), "
           "des cotisations, des prestations versées et du recouvrement des créances.")

    # ── 2. Flux de données métier ─────────────────────────────────────────────
    h1(doc, "2. Flux de Données Métier")
    body(doc,
         "Les données transitent du système source vers les outils de reporting "
         "selon le schéma suivant :")
    add_code_block(doc,
        "Système OLTP (gestion opérationnelle)\n"
        "         ↓  Extraction mensuelle (CLICHE = MMYYYY)\n"
        "   Data Warehouse (DWH)\n"
        "   ├── Tables de dimensions  → référentiels métier stables\n"
        "   └── Tables de faits       → données transactionnelles historisées\n"
        "         ↓  Agrégation\n"
        "   Couche DTM\n"
        "   └── Tables d'indicateurs  → KPIs pré-calculés\n"
        "         ↓\n"
        "   Outils BI (tableaux de bord, rapports)")

    body(doc,
         "Le CLICHE (format MMYYYY, ex : 032026 pour Mars 2026) identifie chaque "
         "snapshot mensuel et permet de comparer les données dans le temps, mois par mois.")

    # ── 3. Domaine GRH ────────────────────────────────────────────────────────
    h1(doc, "3. Domaine GRH – Gestion des Ressources Humaines")
    body(doc,
         "Ce domaine couvre l'ensemble des données relatives au personnel interne "
         "de la CNSS : agents, situations administratives, mouvements, formations, "
         "absences et décisions RH.")

    h2(doc, "3.1 Tables de Dimensions GRH")
    body(doc,
         f"Le domaine GRH comprend {len(GRH_DIMS)} tables de dimensions. "
         "Ces tables contiennent les référentiels métier utilisés pour qualifier "
         "les faits (ex : types de mouvement, fonctions, lieux) :")
    table = doc.add_table(rows=1, cols=3)
    style_table(table)
    add_header_row(table, ["Table cible", "Description métier", "Clé(s) naturelle(s)"])
    add_table_rows(table, GRH_DIMS)

    h2(doc, "3.2 Tables de Faits GRH")
    body(doc,
         f"Le domaine GRH comprend {len(GRH_FACTS)} tables de faits. "
         "Chaque snapshot mensuel (CLICHE) contient l'état complet ou les "
         "transactions de la période :")
    table = doc.add_table(rows=1, cols=2)
    style_table(table)
    add_header_row(table, ["Table cible", "Contenu fonctionnel"])
    add_table_rows(table, GRH_FACTS)

    h2(doc, "3.3 Indicateurs Agrégés GRH (DTM)")
    body(doc,
         "Les indicateurs agrégés sont pré-calculés mensuellement dans la couche DTM "
         "à partir des tables de faits DWH. Ils alimentent directement les tableaux "
         "de bord de pilotage RH.")

    for name, label, desc in GRH_DTM:
        h3(doc, f"{label}  —  {name}")
        body(doc, desc)

    # ── 4. Domaine Prestation / Recouvrement ─────────────────────────────────
    h1(doc, "4. Domaine Prestation / Recouvrement")
    body(doc,
         "Ce domaine couvre la gestion des employeurs affiliés, des travailleurs déclarés, "
         "des cotisations dues et encaissées, des prestations versées (espèces et nature), "
         "et du recouvrement des créances.")

    h2(doc, "4.1 Tables de Dimensions Prestation / Recouvrement")
    body(doc,
         f"Le domaine Prestation/Recouvrement comprend {len(PR_DIMS)} tables de dimensions :")
    table = doc.add_table(rows=1, cols=3)
    style_table(table)
    add_header_row(table, ["Table cible", "Description métier", "Clé(s) naturelle(s)"])
    add_table_rows(table, PR_DIMS)

    h2(doc, "4.2 Tables de Faits Prestation / Recouvrement")
    body(doc,
         f"Le domaine comprend {len(PR_FACTS)} tables de faits couvrant l'ensemble "
         "des activités opérationnelles :")
    table = doc.add_table(rows=1, cols=2)
    style_table(table)
    add_header_row(table, ["Table cible", "Contenu fonctionnel"])
    add_table_rows(table, PR_FACTS)

    h2(doc, "4.3 Indicateurs Agrégés Prestation / Recouvrement (DTM)")
    body(doc,
         f"Le domaine comprend {len(PR_DTM)} tables d'indicateurs agrégés :")

    for name, label, desc in PR_DTM:
        h3(doc, f"{label}  —  {name}")
        body(doc, desc)

    # ── 5. Indicateurs Clés (KPIs) ────────────────────────────────────────────
    h1(doc, "5. Indicateurs Clés (KPIs) Produits")
    table = doc.add_table(rows=1, cols=3)
    style_table(table)
    add_header_row(table, ["Indicateur", "Domaine", "Table DTM source"])
    add_table_rows(table, [
        ("Effectif du personnel par tranche d'âge et ancienneté", "GRH",            "DTM_GRH_EFFECTIF"),
        ("Taux d'encadrement",                                     "GRH",            "DTM_GRH_RATIO"),
        ("Répartition H/F",                                        "GRH",            "DTM_GRH_RATIO"),
        ("Taux de réalisation du plan de formation",               "GRH",            "DTM_GRH_FORMATION"),
        ("Nombre de mouvements entrants / sortants",               "GRH",            "DTM_GRH_MOUVEMENT"),
        ("Taux de recouvrement des cotisations",                   "P/R",            "DTM_COTISATIONS"),
        ("Montant des cotisations appelées vs encaissées",         "P/R",            "DTM_COTISATIONS"),
        ("Nombre d'employeurs affiliés actifs",                    "P/R",            "DTM_EMPLOYEUR"),
        ("Nouvelles immatriculations employeurs / travailleurs",   "P/R",            "DTM_IMM_EMPLOYEUR / DTM_IMM_TRAVAILLEUR"),
        ("Montant des prestations versées par type",               "P/R",            "DTM_PRESTATION"),
        ("Volume et coût des accidents du travail",                "P/R",            "DTM_ACCIDENT_TRAVAIL"),
        ("Nombre et taux de résolution des dossiers",              "P/R",            "DTM_DOSSIER"),
        ("Mises en demeure émises / réglées",                      "P/R",            "DTM_MISE_EN_DEMEURE"),
        ("Prestations indues – montant à récupérer",               "P/R",            "DTM_PRESTATION_INDUE"),
        ("Masse salariale déclarée par secteur",                   "P/R",            "DTM_SALAIRE"),
        ("Taux de dépôt des déclarations de cotisation",           "P/R",            "FAIT_TXDEPOT"),
    ])

    # ── 6. Glossaire ──────────────────────────────────────────────────────────
    h1(doc, "6. Glossaire")
    table = doc.add_table(rows=1, cols=2)
    style_table(table)
    add_header_row(table, ["Terme", "Définition"])
    add_table_rows(table, [
        ("CLICHE",     "Identifiant de snapshot mensuel au format MMYYYY (ex : 032026). "
                       "Permet d'historiser et de rejouer les données mois par mois."),
        ("DWH",        "Data Warehouse – entrepôt de données central hébergeant les tables "
                       "de faits et de dimensions au niveau de détail."),
        ("DTM",        "Data Mart – couche d'agrégats pré-calculés exposés aux outils BI."),
        ("ODS",        "Operational Data Store – schéma d'archivage optionnel. "
                       "Conserve l'historique des données avant TRUNCATE du DWH."),
        ("MERGE",      "Instruction Oracle combinant INSERT et UPDATE : met à jour "
                       "les lignes existantes et insère les nouvelles en une seule opération."),
        ("GTT",        "Global Temporary Table – table temporaire Oracle utilisée comme "
                       "intermédiaire pour les grands MERGE. Réinitialисée à chaque session."),
        ("ETL",        "Extract-Transform-Load – processus d'extraction, transformation "
                       "et chargement des données depuis le système source vers le DWH."),
        ("OLTP",       "Online Transaction Processing – système opérationnel source "
                       "gérant les transactions en temps réel."),
        ("FAIT_",      "Préfixe des tables de faits dans le schéma DWH. "
                       "Contiennent les données transactionnelles détaillées."),
        ("DIM_",       "Préfixe des tables de dimensions. Contiennent les référentiels "
                       "métier utilisés pour qualifier les faits."),
        ("DTM_",       "Préfixe des tables d'indicateurs agrégés dans le schéma DTM."),
        ("PECO_ID",    "Identifiant de la période de cotisation d'un employeur. "
                       "Utilisé dans les tables de cotisation et de transaction."),
        ("AT",         "Accidents du Travail – branche de couverture des accidents "
                       "survenus dans le cadre professionnel."),
        ("AV",         "Assurance Vieillesse – branche de retraite de la sécurité sociale."),
        ("MED",        "Mise en Demeure – acte juridique notifiant à un employeur "
                       "son obligation de s'acquitter de ses cotisations impayées."),
        ("Parquet",    "Format de fichier colonnaire (Apache Parquet) utilisé pour "
                       "le staging intermédiaire. Compressé et efficace pour les grands volumes."),
    ])

    doc.save("documentation_fonctionnelle.docx")
    print("✔  documentation_fonctionnelle.docx généré")


# ─────────────────────────────────────────────────────────────────────────────
# Point d'entrée
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_technical_doc()
    build_functional_doc()
    print("\nDocumentation générée avec succès !")
    print("  → documentation_technique.docx")
    print("  → documentation_fonctionnelle.docx")
